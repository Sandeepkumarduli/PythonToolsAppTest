"""
DOCX to PDF converter tool
"""
import os
import logging
import subprocess
import platform
from docx import Document

def convert_docx_to_pdf(input_path, output_path):
    """
    Convert DOCX file to PDF format
    
    Args:
        input_path (str): Path to input DOCX file
        output_path (str): Path for output PDF file
        
    Returns:
        tuple: (success: bool, message: str)
    """
    try:
        if not os.path.exists(input_path):
            return False, "Input file does not exist"
        
        # Verify the DOCX file can be opened
        try:
            doc = Document(input_path)
            if len(doc.paragraphs) == 0:
                return False, "DOCX file appears to be empty"
        except Exception as e:
            return False, f"Invalid DOCX file: {str(e)}"
        
        # Try different conversion methods based on the platform
        success = False
        error_messages = []
        
        # Method 1: Try using docx2pdf library
        try:
            from docx2pdf import convert
            convert(input_path, output_path)
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                success = True
        except ImportError:
            error_messages.append("docx2pdf library not available")
        except Exception as e:
            error_messages.append(f"docx2pdf conversion failed: {str(e)}")
        
        # Method 2: Try using LibreOffice (if available)
        if not success:
            try:
                # Check if LibreOffice is available
                result = subprocess.run(['libreoffice', '--version'], 
                                      capture_output=True, text=True, timeout=10)
                if result.returncode == 0:
                    # Use LibreOffice for conversion
                    output_dir = os.path.dirname(output_path)
                    result = subprocess.run([
                        'libreoffice', '--headless', '--convert-to', 'pdf',
                        '--outdir', output_dir, input_path
                    ], capture_output=True, text=True, timeout=30)
                    
                    if result.returncode == 0:
                        # LibreOffice creates files with the same name but .pdf extension
                        base_name = os.path.splitext(os.path.basename(input_path))[0]
                        libreoffice_output = os.path.join(output_dir, f"{base_name}.pdf")
                        
                        if os.path.exists(libreoffice_output):
                            if libreoffice_output != output_path:
                                os.rename(libreoffice_output, output_path)
                            success = True
                        else:
                            error_messages.append("LibreOffice conversion completed but output file not found")
                    else:
                        error_messages.append(f"LibreOffice conversion failed: {result.stderr}")
                else:
                    error_messages.append("LibreOffice not available or not working")
            except (subprocess.TimeoutExpired, FileNotFoundError) as e:
                error_messages.append(f"LibreOffice conversion error: {str(e)}")
            except Exception as e:
                error_messages.append(f"LibreOffice conversion unexpected error: {str(e)}")
        
        # Method 3: Create a simple HTML version and try to convert (fallback)
        if not success:
            try:
                # Extract text from DOCX and create a simple HTML file
                doc = Document(input_path)
                html_content = ["<html><head><meta charset='utf-8'></head><body>"]
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        html_content.append(f"<p>{paragraph.text}</p>")
                
                html_content.append("</body></html>")
                
                # Save as HTML temporarily
                html_path = output_path.replace('.pdf', '.html')
                with open(html_path, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(html_content))
                
                # Try to convert HTML to PDF using wkhtmltopdf if available
                try:
                    result = subprocess.run(['wkhtmltopdf', html_path, output_path], 
                                          capture_output=True, text=True, timeout=30)
                    if result.returncode == 0 and os.path.exists(output_path):
                        success = True
                    else:
                        error_messages.append(f"wkhtmltopdf conversion failed: {result.stderr}")
                except (subprocess.TimeoutExpired, FileNotFoundError):
                    error_messages.append("wkhtmltopdf not available")
                
                # Clean up HTML file
                try:
                    os.remove(html_path)
                except:
                    pass
                
            except Exception as e:
                error_messages.append(f"HTML fallback conversion failed: {str(e)}")
        
        if success:
            # Clean up input file
            try:
                os.remove(input_path)
            except Exception as e:
                logging.warning(f"Could not remove input file: {str(e)}")
            
            return True, "DOCX successfully converted to PDF"
        else:
            error_msg = "DOCX to PDF conversion failed. " + \
                       "This feature requires additional system dependencies. " + \
                       "Errors encountered: " + "; ".join(error_messages)
            return False, error_msg
        
    except Exception as e:
        error_msg = f"Error converting DOCX to PDF: {str(e)}"
        logging.error(error_msg)
        return False, error_msg
