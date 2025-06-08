"""
PDF to DOCX converter tool
"""
import os
import logging
from docx import Document
try:
    import PyPDF2
except ImportError:
    try:
        import pypdf as PyPDF2
    except ImportError:
        PyPDF2 = None

def convert_pdf_to_docx(input_path, output_path):
    """
    Convert PDF file to DOCX format
    
    Args:
        input_path (str): Path to input PDF file
        output_path (str): Path for output DOCX file
        
    Returns:
        tuple: (success: bool, message: str)
    """
    try:
        if not PyPDF2:
            return False, "PDF processing library not available. Please install PyPDF2 or pypdf."
        
        if not os.path.exists(input_path):
            return False, "Input file does not exist"
        
        # Create a new Word document
        doc = Document()
        
        # Read PDF and extract text
        with open(input_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            if len(pdf_reader.pages) == 0:
                return False, "PDF file appears to be empty or corrupted"
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():  # Only add non-empty pages
                        if page_num > 0:  # Add page break for subsequent pages
                            doc.add_page_break()
                        
                        # Add page header
                        doc.add_heading(f'Page {page_num + 1}', level=2)
                        
                        # Split text into paragraphs and add to document
                        paragraphs = text.split('\n')
                        for paragraph in paragraphs:
                            if paragraph.strip():
                                doc.add_paragraph(paragraph.strip())
                
                except Exception as e:
                    logging.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                    doc.add_paragraph(f"[Error reading page {page_num + 1}]")
        
        # Save the document
        doc.save(output_path)
        
        # Clean up input file
        try:
            os.remove(input_path)
        except Exception as e:
            logging.warning(f"Could not remove input file: {str(e)}")
        
        return True, "PDF successfully converted to DOCX"
        
    except Exception as e:
        error_msg = f"Error converting PDF to DOCX: {str(e)}"
        logging.error(error_msg)
        return False, error_msg
